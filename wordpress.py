import requests as rq
from bs4 import BeautifulSoup
from bs4.element import Tag, NavigableString
from docx import Document
from docx.shared import Pt
import os

html_path = "HTML"
save_path = "Chapters"


def save_html(html, path):
    with open(path, 'wb') as f:
        f.write(html)


def open_html(path):
    with open(path, 'rb') as f:
        return f.read()


def get_htmls(link, count):
    chapter = rq.get(link).content
    save_html(chapter, f"{html_path}/{count:04}.html")

    soup = BeautifulSoup(chapter, 'html.parser')
    next_link = soup.find(class_="nav-next").contents[0]['href']
    if next_link is not None:
        get_htmls(next_link, count+1)


def write_chapter(html, book_path):

    # setting up text and document
    soup = BeautifulSoup(open_html(html), 'html.parser')
    title = soup.find(class_="entry-title").get_text()
    contents = soup.find(class_="entry-content").find_all('p')
    chapter = Document(book_path)

    # Writing text begins here

    # Title
    title_run = chapter.add_paragraph().add_run(title)
    title_run.bold = True
    title_run.font.size = Pt(36)
    title_run.font.name = 'Noto Sans'

    # Chapter Text
    for section in contents:
        # set up paragraph
        graf = chapter.add_paragraph()
        graf_format = graf.paragraph_format
        graf_format.line_spacing = Pt(20)
        graf_format.space_before = Pt(20)
        graf_format.space_after = Pt(20)

        # add and format runs
        def add_element(ele):
            if isinstance(ele, NavigableString):  # simply add the string
                graf_run = graf.add_run(ele)
                graf_run.font.name = 'Noto Sans'
                return graf_run
            elif ele.name in ('i', 'em'):
                run = add_element(ele.contents[0])
                if run is not None:
                    run.font.italic = True
            elif ele.name in ('b', 'strong'):
                run = add_element(ele.contents[0])
                if run is not None:
                    run.font.bold = True
            else:
                [add_element(child) for child in list(ele.children)]

        add_element(section)

    # End formatting and save
    chapter.add_page_break()
    chapter.save(book_path)


def write_book(book_num, start_ch, end_ch):
    book = Document()
    book.save(f"PTGE_{book_num}.docx")
    for num in range(start_ch, end_ch+1):
        write_chapter(f"HTML/{num:04}.html", f"PTGE_{book_num}.docx")


#write_book(1, 0, 29)
write_book(2, 30, 97)
write_book(3, 98, 208)
write_book(4, 209, 324)
write_book(5, 325, 394)




